"""Parse roster documents (DOCX, XLSX, ZIP) into personnel records with photos."""

from __future__ import annotations

import base64
import csv
import io
import json
import logging
import re
import zipfile
from dataclasses import dataclass
from typing import Any

logger = logging.getLogger(__name__)

VALID_ROLES = {"guard", "vip", "staff", "contractor"}

NAME_HEADERS = ("name", "full name", "person", "employee")
DESIGNATION_HEADERS = ("designation", "title", "position", "job title", "department")
ROLE_HEADERS = ("role", "type", "personnel role", "category")
PHOTO_HEADERS = ("photo", "image", "portrait", "picture", "filename", "photo file")

# Default csv limit is 128KB — inline base64 portraits exceed that.
_CSV_FIELD_LIMIT = 25 * 1024 * 1024


def _ensure_csv_field_limit() -> None:
    try:
        if csv.field_size_limit() < _CSV_FIELD_LIMIT:
            csv.field_size_limit(_CSV_FIELD_LIMIT)
    except OverflowError:
        csv.field_size_limit(min(_CSV_FIELD_LIMIT, 2**31 - 1))


_ensure_csv_field_limit()


@dataclass
class ParsedPerson:
    name: str
    designation: str
    role: str
    photo_bytes: bytes
    photo_mime: str
    row_index: int


class RosterParseError(ValueError):
    """Raised when the roster file cannot be parsed."""


def _normalize_header(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip().lower())


def _find_column(headers: list[str], candidates: tuple[str, ...]) -> int | None:
    normalized = [_normalize_header(h) for h in headers]
    for i, header in enumerate(normalized):
        if header in candidates:
            return i
    for i, header in enumerate(normalized):
        if any(c in header for c in candidates):
            return i
    return None


def _normalize_role(value: str | None, default: str = "staff") -> str:
    raw = (value or default).strip().lower()
    aliases = {
        "security": "guard",
        "security guard": "guard",
        "officer": "guard",
        "vip guest": "vip",
        "guest": "vip",
        "employee": "staff",
        "team": "staff",
        "vendor": "contractor",
    }
    role = aliases.get(raw, raw)
    return role if role in VALID_ROLES else default


def _guess_mime(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".png"):
        return "image/png"
    if lower.endswith(".webp"):
        return "image/webp"
    if lower.endswith(".gif"):
        return "image/gif"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if data[:2] == b"\xff\xd8":
        return "image/jpeg"
    return "image/jpeg"


def _data_url_from_bytes(photo_bytes: bytes, mime: str) -> str:
    encoded = base64.b64encode(photo_bytes).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _decode_photo_field(
    photo_ref: str, image_map: dict[str, bytes]
) -> tuple[bytes, str] | None:
    """Resolve a photo cell: inline base64, data URL, http(s) URL, or ZIP filename."""
    photo_ref = (photo_ref or "").strip()
    if not photo_ref:
        return None

    if photo_ref.startswith("data:image"):
        try:
            header, payload = photo_ref.split(",", 1)
            mime = header.split(";")[0].replace("data:", "")
            return base64.b64decode(payload), mime
        except Exception:
            return None

    if photo_ref.startswith(("http://", "https://")):
        try:
            import httpx

            resp = httpx.get(photo_ref, timeout=15, follow_redirects=True)
            resp.raise_for_status()
            data = resp.content
            content_type = resp.headers.get("content-type", "image/jpeg").split(";")[0]
            mime = content_type if content_type.startswith("image/") else _guess_mime("photo.jpg", data)
            return data, mime
        except Exception as exc:
            logger.warning("Failed to fetch roster photo URL: %s", exc)
            return None

    key = photo_ref.replace("\\", "/").split("/")[-1].lower()
    for fname, data in image_map.items():
        if fname.lower() == key or fname.lower().endswith("/" + key):
            return data, _guess_mime(fname, data)

    compact = re.sub(r"\s+", "", photo_ref)
    if len(compact) > 80 and re.fullmatch(r"[A-Za-z0-9+/=]+", compact):
        try:
            data = base64.b64decode(compact, validate=True)
            if len(data) > 50:
                return data, _guess_mime("photo.jpg", data)
        except Exception:
            pass

    return None


def _cell_image_bytes(cell, document) -> bytes | None:
    from docx.oxml.ns import qn

    for blip in cell._element.findall(".//" + qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if not embed:
            continue
        part = document.part.related_parts.get(embed)
        if part is not None and hasattr(part, "blob"):
            return part.blob
    return None


def _parse_docx_table(document, table) -> list[ParsedPerson]:
    if len(table.rows) < 2:
        return []

    headers = [cell.text.strip() for cell in table.rows[0].cells]
    name_idx = _find_column(headers, NAME_HEADERS)
    if name_idx is None:
        return []

    desig_idx = _find_column(headers, DESIGNATION_HEADERS)
    role_idx = _find_column(headers, ROLE_HEADERS)

    media_queue: list[bytes] = []
    docx_bytes = getattr(document, "_roster_bytes", None)
    if docx_bytes:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            for name in sorted(zf.namelist()):
                if name.startswith("word/media/"):
                    media_queue.append(zf.read(name))

    people: list[ParsedPerson] = []
    for row_num, row in enumerate(table.rows[1:], start=2):
        cells = row.cells
        name = cells[name_idx].text.strip() if name_idx < len(cells) else ""
        if not name:
            continue

        designation = ""
        if desig_idx is not None and desig_idx < len(cells):
            designation = cells[desig_idx].text.strip()
        if not designation:
            designation = "Personnel"

        role = "staff"
        if role_idx is not None and role_idx < len(cells):
            role = _normalize_role(cells[role_idx].text)

        photo_bytes: bytes | None = None
        for cell in cells:
            photo_bytes = _cell_image_bytes(cell, document)
            if photo_bytes:
                break

        if photo_bytes is None and media_queue:
            photo_bytes = media_queue.pop(0)

        if not photo_bytes:
            logger.warning("Row %s (%s): no photo found", row_num, name)
            continue

        people.append(
            ParsedPerson(
                name=name,
                designation=designation,
                role=role,
                photo_bytes=photo_bytes,
                photo_mime=_guess_mime(f"row{row_num}.jpg", photo_bytes),
                row_index=row_num,
            )
        )

    return people


def parse_docx(content: bytes) -> list[ParsedPerson]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RosterParseError(
            "DOCX support is not installed. Run: pip install python-docx"
        ) from exc

    document = Document(io.BytesIO(content))
    document._roster_bytes = content  # type: ignore[attr-defined]

    people: list[ParsedPerson] = []
    for table in document.tables:
        people.extend(_parse_docx_table(document, table))

    if people:
        return people

    # Paragraph fallback: "Name — Designation" lines + sequential images
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    media: list[bytes] = []
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        for name in sorted(zf.namelist()):
            if name.startswith("word/media/"):
                media.append(zf.read(name))

    for i, line in enumerate(lines):
        if "—" in line:
            name, designation = [part.strip() for part in line.split("—", 1)]
        elif " - " in line:
            name, designation = [part.strip() for part in line.split(" - ", 1)]
        else:
            name, designation = line, "Personnel"

        if not name or i >= len(media):
            continue
        people.append(
            ParsedPerson(
                name=name,
                designation=designation or "Personnel",
                role="staff",
                photo_bytes=media[i],
                photo_mime=_guess_mime(f"line{i}.jpg", media[i]),
                row_index=i + 1,
            )
        )

    if not people:
        raise RosterParseError(
            "No roster found in document. Use a table with columns: Name, Designation, Role, Photo."
        )
    return people


def _parse_csv_rows(rows: list[dict[str, str]], image_map: dict[str, bytes]) -> list[ParsedPerson]:
    people: list[ParsedPerson] = []
    for i, row in enumerate(rows, start=1):
        lowered = {_normalize_header(k): (v or "").strip() for k, v in row.items()}

        name = ""
        for key in NAME_HEADERS:
            if lowered.get(key):
                name = lowered[key]
                break
        if not name:
            for val in lowered.values():
                if val and not any(val.lower().endswith(ext) for ext in (".jpg", ".png", ".jpeg", ".webp")):
                    name = val
                    break

        if not name:
            continue

        designation = ""
        for key in DESIGNATION_HEADERS:
            if lowered.get(key):
                designation = lowered[key]
                break
        if not designation:
            designation = "Personnel"

        role = "staff"
        for key in ROLE_HEADERS:
            if lowered.get(key):
                role = _normalize_role(lowered[key])
                break

        photo_ref = ""
        for key in PHOTO_HEADERS:
            if lowered.get(key):
                photo_ref = lowered[key]
                break

        decoded = _decode_photo_field(photo_ref, image_map) if photo_ref else None
        photo_bytes: bytes | None = None
        mime = "image/jpeg"
        if decoded:
            photo_bytes, mime = decoded

        if not photo_bytes and len(image_map) == len(rows):
            ordered = list(image_map.values())
            photo_bytes = ordered[i - 1]
            mime = _guess_mime(f"row{i}.jpg", photo_bytes)

        if not photo_bytes:
            continue

        people.append(
            ParsedPerson(
                name=name,
                designation=designation,
                role=role,
                photo_bytes=photo_bytes,
                photo_mime=mime,
                row_index=i,
            )
        )
    return people


def parse_zip_roster(content: bytes) -> list[ParsedPerson]:
    image_map: dict[str, bytes] = {}
    manifest_text: str | None = None
    manifest_name: str | None = None
    json_rows: list[dict[str, Any]] | None = None

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = info.filename.replace("\\", "/")
            lower = name.lower()
            data = zf.read(info)

            if lower.endswith((".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp")):
                image_map[name.split("/")[-1]] = data
                image_map[name.lower()] = data
                continue

            if lower.endswith(".json") and ("roster" in lower or "manifest" in lower or lower.endswith("/people.json")):
                try:
                    payload = json.loads(data.decode("utf-8"))
                    if isinstance(payload, list):
                        json_rows = payload
                    elif isinstance(payload, dict) and isinstance(payload.get("personnel"), list):
                        json_rows = payload["personnel"]
                except json.JSONDecodeError:
                    pass
                continue

            if lower.endswith(".csv") and manifest_text is None:
                manifest_text = data.decode("utf-8-sig")
                manifest_name = name

    if json_rows:
        people: list[ParsedPerson] = []
        for i, row in enumerate(json_rows, start=1):
            if not isinstance(row, dict):
                continue
            name = str(row.get("name") or "").strip()
            if not name:
                continue
            designation = str(row.get("designation") or row.get("title") or "Personnel").strip()
            role = _normalize_role(str(row.get("role") or "staff"))
            photo_bytes: bytes | None = None
            mime = "image/jpeg"

            photo_field = row.get("photo") or row.get("photoUrl") or row.get("image")
            if isinstance(photo_field, str):
                if photo_field.startswith("data:image"):
                    try:
                        header, payload = photo_field.split(",", 1)
                        mime = header.split(";")[0].replace("data:", "")
                        photo_bytes = base64.b64decode(payload)
                    except Exception:
                        photo_bytes = None
                else:
                    key = photo_field.replace("\\", "/").split("/")[-1].lower()
                    photo_bytes = image_map.get(key) or image_map.get(photo_field.lower())

            if not photo_bytes:
                continue

            people.append(
                ParsedPerson(
                    name=name,
                    designation=designation,
                    role=role,
                    photo_bytes=photo_bytes,
                    photo_mime=mime,
                    row_index=i,
                )
            )
        if people:
            return people

    if manifest_text:
        _ensure_csv_field_limit()
        reader = csv.DictReader(io.StringIO(manifest_text))
        rows = list(reader)
        people = _parse_csv_rows(rows, image_map)
        if people:
            return people

    if image_map:
        people = []
        for i, (fname, data) in enumerate(sorted(image_map.items()), start=1):
            stem = fname.rsplit(".", 1)[0]
            name = stem.replace("_", " ").replace("-", " ").title()
            people.append(
                ParsedPerson(
                    name=name,
                    designation="Personnel",
                    role="staff",
                    photo_bytes=data,
                    photo_mime=_guess_mime(fname, data),
                    row_index=i,
                )
            )
        return people

    raise RosterParseError(
        "ZIP must contain roster.csv (or roster.json) plus portrait images, or image files named by person."
    )


def _extract_xlsx_zip_media(content: bytes) -> list[tuple[bytes, str]]:
    """Images stored in xl/media/ (e.g. Excel mobile exports without cell anchors)."""
    out: list[tuple[bytes, str]] = []
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        for name in sorted(zf.namelist()):
            if name.startswith("xl/media/"):
                data = zf.read(name)
                out.append((data, _guess_mime(name, data)))
    return out


def parse_xlsx(content: bytes) -> list[ParsedPerson]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RosterParseError(
            "Excel support is not installed. Run: pip install openpyxl"
        ) from exc

    wb = load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active
    if ws is None or ws.max_row < 2:
        raise RosterParseError("Excel sheet is empty.")

    headers = [str(ws.cell(1, c).value or "").strip() for c in range(1, ws.max_column + 1)]
    name_col = _find_column(headers, NAME_HEADERS)
    if name_col is None:
        raise RosterParseError("Excel must have a Name column in the first row.")

    desig_col = _find_column(headers, DESIGNATION_HEADERS)
    role_col = _find_column(headers, ROLE_HEADERS)
    photo_col = _find_column(headers, PHOTO_HEADERS)

    images_by_row: dict[int, tuple[bytes, str]] = {}
    for image in getattr(ws, "_images", []):
        try:
            anchor = image.anchor._from  # type: ignore[attr-defined]
            row = int(anchor.row) + 1
            col = int(anchor.col) + 1
            data = image._data()  # type: ignore[attr-defined]
            fmt = getattr(image, "format", "jpeg") or "jpeg"
            mime = f"image/{fmt.lower()}"
            images_by_row[row] = (data, mime)
            images_by_row[col] = (data, mime)
        except Exception:
            logger.debug("Skipped embedded Excel image", exc_info=True)

    zip_media = _extract_xlsx_zip_media(content)
    zip_only_mode = not images_by_row and bool(zip_media)
    data_row_index = 0

    people: list[ParsedPerson] = []
    for row in range(2, ws.max_row + 1):
        name = str(ws.cell(row, name_col + 1).value or "").strip()
        if not name:
            continue

        designation = "Personnel"
        if desig_col is not None:
            designation = str(ws.cell(row, desig_col + 1).value or "").strip() or designation

        role = "staff"
        if role_col is not None:
            role = _normalize_role(str(ws.cell(row, role_col + 1).value or ""))

        photo_bytes: bytes | None = None
        mime = "image/jpeg"
        if row in images_by_row:
            photo_bytes, mime = images_by_row[row]
        elif photo_col is not None:
            val = str(ws.cell(row, photo_col + 1).value or "").strip()
            if val and not val.startswith("#"):
                decoded = _decode_photo_field(val, {})
                if decoded:
                    photo_bytes, mime = decoded

        if photo_bytes is None and zip_only_mode and data_row_index < len(zip_media):
            photo_bytes, mime = zip_media[data_row_index]

        data_row_index += 1

        if not photo_bytes:
            continue

        people.append(
            ParsedPerson(
                name=name,
                designation=designation,
                role=role,
                photo_bytes=photo_bytes,
                photo_mime=mime,
                row_index=row,
            )
        )

    if not people:
        raise RosterParseError(
            "No enrollable rows in this Excel file. Embed portrait photos inside the sheet "
            "(Insert → Pictures in each row), or export as CSV UTF-8 with base64 in a photo column, "
            "or ZIP roster.csv plus image files."
        )
    return people


def parse_roster_file(filename: str, content: bytes) -> list[ParsedPerson]:
    if not content:
        raise RosterParseError("Uploaded file is empty.")

    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx(content)
    if lower.endswith(".zip"):
        return parse_zip_roster(content)
    if lower.endswith((".xlsx", ".xlsm")):
        return parse_xlsx(content)
    if lower.endswith(".csv"):
        if content[:2] == b"PK":
            raise RosterParseError(
                "This file is Excel format, not CSV. In Excel use Save As → CSV UTF-8, "
                "or upload the .xlsx with portrait photos embedded in the sheet cells."
            )
        _ensure_csv_field_limit()
        text = content.decode("utf-8-sig")
        rows = list(csv.DictReader(io.StringIO(text)))
        people = _parse_csv_rows(rows, {})
        if not people:
            filename_refs = 0
            missing_photo_col = True
            for row in rows:
                lowered = {_normalize_header(k): (v or "").strip() for k, v in row.items()}
                for key in PHOTO_HEADERS:
                    ref = lowered.get(key, "")
                    if ref:
                        missing_photo_col = False
                        if ref.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
                            filename_refs += 1
            if filename_refs:
                raise RosterParseError(
                    "CSV lists photo filenames (e.g. john.jpg) but not the image data. "
                    "ZIP roster.csv with the portrait files, or paste base64 / data:image URLs "
                    "into the photo column."
                )
            if missing_photo_col:
                raise RosterParseError(
                    "CSV needs columns: name, designation, role, photo. "
                    "Put each portrait as base64 or data:image/jpeg;base64,... in the photo cell."
                )
            raise RosterParseError(
                "No rows could be enrolled from this CSV. Each row needs a readable portrait "
                "in the photo column (base64 or data URL), and a clear front-facing face."
            )
        return people

    raise RosterParseError(
        "Unsupported file type. Upload .csv, .docx, .xlsx, or .zip roster."
    )
